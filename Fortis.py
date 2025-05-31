import streamlit as st
import openai
import os
from datetime import datetime
import pandas as pd
import docx
from docx import Document
import io
import re

# Set your assistant ID here
ASSISTANT_ID = "asst_vr6ZFdxJVO4kd4oukqMUTk9m"

# Load API key securely from Streamlit secrets
openai.api_key = st.secrets["OPENAI_API_KEY"]

st.set_page_config(page_title="Financial Report Generator", layout="centered")
st.title("📊 Financial Report Generator")

st.markdown("Upload your client data as a '.txt', '.csv', or '.docx' file. The assistant will generate a detailed financial advice report based on the contents.")

uploaded_file = st.file_uploader("Upload a file", type=["txt", "csv", "docx"])

def extract_text(file):
    file_type = file.name.split(".")[-1].lower()

    if file_type == "txt":
        return file.read().decode("utf-8")

    elif file_type == "csv":
        df = pd.read_csv(file)
        return df.to_string(index=False)

    elif file_type == "docx":
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])

    else:
        return None

if uploaded_file:
    # If this is a new upload (different file name), clear the old “generated” flag
    if st.session_state.get("last_upload") != uploaded_file.name:
        st.session_state["last_upload"]      = uploaded_file.name
        st.session_state.pop("report_generated", None)

    # 1️⃣ Extract client input once
    client_input = extract_text(uploaded_file)
    if not client_input:
        st.error("Unsupported file format or failed to extract text.")
        st.stop()

    # 2️⃣ If we haven’t generated yet, show a Generate button
    if "report_generated" not in st.session_state:
        if st.button("Generate Report"):
            with st.spinner("Generating report…"):
                # ← Paste your entire AI + regex + DOCX-building logic here —
                #     everything from today_date up to storing in session_state. For example:

                # Inject today's date
                today_date = datetime.now().strftime("%d %B %Y")
                system_injected_text = f"Today's date is {today_date}.\n\n{client_input}"

                # OpenAI thread/run logic...
                thread = openai.beta.threads.create()
                openai.beta.threads.messages.create(
                    thread_id=thread.id,
                    role="user",
                    content=system_injected_text
                )
                run = openai.beta.threads.runs.create(
                    thread_id=thread.id,
                    assistant_id=ASSISTANT_ID
                )
                # wait for completion...
                while True:
                    status = openai.beta.threads.runs.retrieve(
                        thread_id=thread.id, run_id=run.id
                    )
                    if status.status == "completed":
                        break
                    elif status.status == "failed":
                        st.error("Report generation failed. Please try again.")
                        st.stop()

                messages = openai.beta.threads.messages.list(thread_id=thread.id)
                report_text = messages.data[0].content[0].text.value

                # Markdown normalization (your regex)...
                report_text = re.sub(r"^\*\*(.+?)\*\*$", r"## \1",
                                     report_text, flags=re.MULTILINE)
                report_text = re.sub(r"^(## .+?)\s*\n+", r"\1\n\n",
                                     report_text, flags=re.MULTILINE)

                # Build DOCX in memory...
                doc = Document()
                bold_pattern = re.compile(r"\*\*(.+?)\*\*")
                for line in report_text.splitlines():
                    p = doc.add_paragraph(); last_end = 0
                    for m in bold_pattern.finditer(line):
                        if m.start() > last_end:
                            p.add_run(line[last_end:m.start()])
                        run = p.add_run(m.group(1)); run.bold = True
                        last_end = m.end()
                    if last_end < len(line):
                        p.add_run(line[last_end:])
                word_file = io.BytesIO(); doc.save(word_file); word_file.seek(0)

                # Store in session_state
                st.session_state.report_text      = report_text
                st.session_state.word_file        = word_file
                st.session_state.report_generated = True

    # 3️⃣ Once generated (or on repeat run), display & download
    if st.session_state.get("report_generated"):
        st.subheader("📄 Generated Report")
        st.markdown(st.session_state.report_text)
        st.download_button(
            "📥 Download Report (.docx)",
            data=st.session_state.word_file,
            file_name=f"financial_report_{datetime.now():%Y%m%d_%H%M%S}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
